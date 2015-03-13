# -*- coding: utf-8 -*-
'''
Created : 2015-03-12

@author: Eric Lapouyade
'''

__version__ = '0.1.1'

from lxml import etree
from docx import Document
from jinja2 import Template
import copy
import re

class DocxTemplate(object):
    """ Class for managing docx files as they were jinja2 templates """
    def __init__(self, docx):
        self.docx = Document(docx)
    
    def __getattr__(self, name):        
        return getattr(self.docx, name)
    
    def get_docx(self):
        return self.docx
    
    def get_xml(self):
        return etree.tostring(self.docx._element.body, pretty_print=True)

    def write_xml(self,filename):
        with open(filename,'w') as fh:
            fh.write(self.get_xml())
    
    def patch_xml(self,src_xml):
        # strip all xml tags inside {% %} and {{ }}
        # that Microsoft word can insert into xml code for this part of the document
        def striptags(m):
            return re.sub('</w:t>.*?(<w:t>|<w:t [^>]*>)','',m.group(0),flags=re.DOTALL)
        src_xml = re.sub(r'{%(?:(?!%}).)*|{{(?:(?!}}).)*',striptags,src_xml,flags=re.DOTALL)
        
        # manage table cell background color
        def cellbg(m):
            cell_xml = m.group(1) + m.group(3)
            cell_xml = re.sub(r'<w:r[ >](?:(?!<w:r[ >]).)*<w:t></w:t>.*?</w:r>','',cell_xml,flags=re.DOTALL)
            cell_xml = re.sub(r'<w:shd[^/]*/>','', cell_xml, count=1)
            return re.sub(r'(<w:tcPr[^>]*>)',r'\1<w:shd w:val="clear" w:color="auto" w:fill="{{%s}}"/>' % m.group(2), cell_xml)
        src_xml = re.sub(r'(<w:tc[ >](?:(?!<w:tc[ >]).)*){%\s*cellbg\s+([^%]*)\s*%}(.*?</w:tc>)',cellbg,src_xml,flags=re.DOTALL)
        
        for y in ['tr', 'p', 'r']:
            # replace into xml code the row containing {%y xxx %} or {{y xxx}} template tag 
            # by {% xxx %} or {{ xx }} without any surronding xml tags.
            pat = r'<w:%(y)s[ >](?:(?!<w:%(y)s[ >]).)*({%%|{{)%(y)s ([^}%%]*(?:%%}|}})).*?</w:%(y)s>' % {'y':y}
            src_xml = re.sub(pat, r'\1 \2',src_xml,flags=re.DOTALL)
        
        return src_xml

    def render_xml(self,src_xml,context):
        template = Template(src_xml)
        dst_xml = template.render(context)    
        return dst_xml

    def build_xml(self,context):
        xml = self.get_xml()
        xml = self.patch_xml(xml)
        xml = self.render_xml(xml, context)
        return xml
        
    def map_xml(self,xml):
        root = self.docx._element
        body = root.body
        root.replace(body,etree.fromstring(xml))

    def render(self,context):
        xml = self.build_xml(context)
        self.map_xml(xml)
        
    def new_subdoc(self):
        return Subdoc(self)

class Subdoc(object):
    """ Class for subdocumentation insertion into master document """
    def __init__(self, tpl):
        self.tpl = tpl
        self.docx = tpl.get_docx()
        self.subdocx = Document()
        self.subdocx._part = self.docx._part

    def __getattr__(self, name):        
        return getattr(self.subdocx, name)
    
    def __unicode__(self):
        xml = ''
        for p in self.paragraphs:
            xml += '<w:p>\n' + re.sub(r'^.*\n', '', etree.tostring(p._element,pretty_print=True))
        return xml
    
    

class Tpldoc(object):
    """ class to build documenation to be passed into template variables """
    pass